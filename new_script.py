import funtions as fp
import time
from scipy import stats
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches
import numpy as np
import matplotlib.pyplot as plt
tabla_frecuenciaindex=['marca','Freq','FrecAbs','Frecrel','Frecrelacum']
estadisticosindex=['media','Erro estandar media','moda','mediana','desviacion_estandar','varianza']
estadisticosindex2=["Q1","Q3","rango intercuartil","Asimetria","kurtosis"]
tesnormalidad=["variable","valor D","Valor P","conclusion"]
intervaloconfianza2=["variable","intervalo inferior","intervalo superior"]
proyecto=Document()
#proyecto.add_heading('Escuela superior politecnica del litoral',level=0)
proyecto.add_heading("Introduccion",level=1)
proyecto.add_heading("\n")
proyecto.add_heading("Objetivos",level=1)
proyecto.add_heading("\n")
proyecto.add_heading("objetivo general",level=2)
proyecto.add_heading("\n")
proyecto.add_heading("Objetivos especificos",level=2)
proyecto.add_heading("\n")
proyecto.add_heading("Variables y Metodología",level=1)
proyecto.add_heading("\n")
proyecto.add_heading("Análisis Descriptivo de Datos",level=1)
proyecto.add_heading("\n")

variables,tipo,df=fp.abrir_archivo("divorcios.xlsx")

contador=0
plt.style.use('ggplot')
for i in variables:

        if str(tipo[contador]).count("int") > 0 or str(tipo[contador]).count("float") > 0:
            proyecto.add_heading("Analisis Descriptivo para la variable cuantitativa {}".format(i), level=3)
            proyecto.add_heading("\n")
            valor3=np.array(df[i].tolist())
            tabla,conj,his=fp.tabla_fercuencia(df[i].tolist())
            proyecto.add_paragraph('estadisticos descriptivos')
            proyecto.add_heading("\n")
            """
              |
              | tabla 1
              |
            """
            tabla1=proyecto.add_table(rows=1,cols=len(estadisticosindex))
            hdr_cells = tabla1.rows[0].cells
            for j in range(len(estadisticosindex)):
                hdr_cells[j].text = estadisticosindex[j]
            row_cells = tabla1.add_row().cells
            row_cells[0].text = str(round(valor3.mean(), 1))
            row_cells[1].text = str(round(valor3.std() / np.sqrt(len(valor3)), 1))
            row_cells[2].text = str(stats.mode(valor3)[0][0])
            row_cells[3].text = str(round(np.percentile(np.array(valor3), 50), 1))
            row_cells[4].text = str(round(valor3.std(),1))
            row_cells[5].text = str(round(valor3.var(),1))
            proyecto.add_paragraph("\n")
            """
             |
             |tabla 2
            """
            tabla2=proyecto.add_table(rows=1, cols=len(estadisticosindex2))
            hdr_cells2 = tabla2.rows[0].cells
            for j in range(len(estadisticosindex2)):
                hdr_cells2[j].text = estadisticosindex2[j]
            row_cells2=tabla2.add_row().cells
            row_cells2[0].text = str(round(np.percentile(valor3,25),2))
            row_cells2[1].text = str(round(np.percentile(valor3,75),2))
            row_cells2[2].text = str(round(np.percentile(valor3,75)-np.percentile(valor3,25),2))
            row_cells2[3].text = str(round(stats.skew(valor3),2))
            row_cells2[4].text = str(round(stats.kurtosis(valor3),2))
            proyecto.add_paragraph("\n")
            """
             |
             | tabla 3
             |
            """
            proyecto.add_paragraph('tabla de frecuencias')
            proyecto.add_heading("\n")
            tabla3=proyecto.add_table(rows=1, cols=len(tabla_frecuenciaindex))
            hdr_cells3=tabla3.rows[0].cells
            for j in range(len(tabla_frecuenciaindex)):
                hdr_cells3[j].text = tabla_frecuenciaindex[j]
            for marca,Freq, FrecAbs, Frecrel, Frecrelacumin in tabla:
                row_cells3 = tabla3.add_row().cells
                row_cells3[0].text = str(marca)
                row_cells3[1].text = str(round(Freq,2))
                row_cells3[2].text = str(round(FrecAbs,2))
                row_cells3[3].text = str(round(Frecrel,2))
                row_cells3[4].text = str(round(Frecrelacumin,2))
            proyecto.add_paragraph("\n")
            """
             |
             |graficos
            """
            ecdf=sm.distributions.ECDF(valor3)
            x = np.linspace(np.min(valor3),np.max(valor3))
            y=ecdf(x)
            proyecto.add_paragraph('Analisis grafico')
            fig,axs=plt.subplots(2,2,gridspec_kw={'hspace':0.4,"wspace":0.25})
            fig.suptitle("Analisis grafico de la variable {}".format(i))
            axs[0,0].hist(valor3,bins=conj,density=True)
            axs[0,0].set_title("histograma")
            axs[0,0].set(ylabel="densidad")
            axs[0,1].boxplot(valor3,sym='ko',whis=1.5)
            axs[0,1].set_title("Diagrama de cajas")
            #axs[0,1].set(xlabel="intervalos", ylabel="densidad")
            axs[1,0].plot(conj[:len(conj)-1],his.cumsum(),"o-")
            axs[1,0].set_title("Grafico de ojiva")
            axs[1,0].set(xlabel="observaciones", ylabel="densidad acumulada")
            axs[1,1].step(x,y)
            axs[1,1].set_title("distribucion empirica")
            nombre="grafico_variable_{}.png".format(i)
            plt.savefig(nombre)
            #time.sleep(1)
            proyecto.add_picture(nombre,width=Inches(6))
            """
             |parrafo
             |
            """
            proyecto.add_paragraph(fp.escritura(i,var1=round(stats.skew(valor3),2),var2=round(stats.kurtosis(valor3),2),var3=round(np.percentile(valor3,50),2),var4=round(np.percentile(valor3,75),2),var5=round(np.percentile(valor3,25),2)))
            proyecto.add_page_break()

        else:
            proyecto.add_heading("Analisis Descriptivo para la variable cualitativa {}".format(i), level=3)
            proyecto.add_heading("\n")
            tablacualiindex = ['Valores','Frecuencia']
            tablacuali = []
            valores = []
            valor4=list(set(df[i].tolist()))
            for j in valor4:
                conteo_var = list(df[i].tolist()).count(j)
                valores.append(list(df[i].tolist()).count(j))
                tuplita_2 =(j,conteo_var)
                tablacuali.append(tuplita_2)
            tablacuali = tuple(tablacuali)
            """
             |tabla 4
             |
            """
            tabla4 =proyecto.add_table(rows=1, cols=2)
            hdr_cells4= tabla4.rows[0].cells
            for j in range(len(tablacualiindex)):
                hdr_cells4[j].text = tablacualiindex[j]
            for j in tablacuali:
                row_cells4= tabla4.add_row().cells
                row_cells4[0].text = str(j[0])
                row_cells4[1].text = str(j[1])
            proyecto.add_paragraph("\n")
            plt.figure()
            plt.bar(valor4,valores)
            plt.xticks(rotation=20)
            plt.suptitle('graficó de barras para la variable cualitativa {}'.format(i))
            imagenabrir2 = 'analisis_var_{}.png'.format(i)
            plt.savefig('analisis_var_{}.png'.format(i))
            proyecto.add_picture(imagenabrir2,width=Inches(4))
            proyecto.add_page_break()
        contador+=1

proyecto.add_heading("Análisis Estadístico Bivariado",level=1)
proyecto.add_heading("\n")

proyecto.add_heading("Estadística Inferencial",level=1)
proyecto.add_heading("\n")
proyecto.add_heading("Bondad de ajuste",level=2)
proyecto.add_heading("\n")
lista=[]
lista3=[]
contador=0
for i in variables:
     if str(tipo[contador]).count("int")>0 or str(tipo[contador]).count("float")>0:
          #print(tabla_fercuencia(df[i].tolist()))
          lista.append(fp.kolmogorovtest(i,df[i].tolist()))
          lista3.append(fp.intervaloconfianza(i,df[i].tolist()))
     contador+=1

lista=tuple(lista)
lista3=tuple(lista3)
"""
 |tabla 5
 |
"""
tabla5 =proyecto.add_table(rows=1,cols=len(tesnormalidad))
hdr_cells5= tabla5.rows[0].cells
for j in range(len(tesnormalidad)):
    hdr_cells5[j].text = tesnormalidad[j]
for var,vard,varp,rel in lista:
    row_cells5=tabla5.add_row().cells
    row_cells5[0].text = str(var)
    row_cells5[1].text = str(round(vard,8))
    row_cells5[2].text=str(round(varp,4))
    row_cells5[3].text=str(rel)
proyecto.add_paragraph("\n")

proyecto.add_heading("Intervalo de confianza",level=2)
proyecto.add_heading("\n")
"""
 |tabla 6
 |
"""
tabla6=proyecto.add_table(rows=1,cols=len(intervaloconfianza2))
hdr_cells6= tabla6.rows[0].cells
for j in range(len(intervaloconfianza2)):
    hdr_cells6[j].text=intervaloconfianza2[j]
for j in lista3:
    row_cells6=tabla6.add_row().cells
    row_cells6[0].text = str(j[0])
    row_cells6[1].text = str(round(j[1],4))
    row_cells6[2].text=str(round(j[2],4))
proyecto.add_heading("\n")

proyecto.add_heading("Analisis de contingencia",level=2)
proyecto.add_heading("\n")
proyecto.add_heading("Pruebas de hipotesis",level=2)
proyecto.add_heading("\n")
proyecto.add_heading("Conclusiones",level=1)
proyecto.add_heading("\n")
proyecto.save('divorcios.docx')


